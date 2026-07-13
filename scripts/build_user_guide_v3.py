"""Build the user guide v3 (.docx) — compact, ~5 pages.

    python scripts/build_user_guide_v3.py                       # default: adviser
    python scripts/build_user_guide_v3.py --edition adviser     # 4-tab variant
    python scripts/build_user_guide_v3.py --edition full        # 5-tab variant (CLASS Import included)

Adviser edition: omits CLASS Import staging tab. For advisers who don't use
CLASS Super as administration software.

Full edition: includes CLASS Import staging tab. For advisers using CLASS
Super — adds a Quick Start branch and a CLASS workflow mini-section to
the Inputs reference.

Layout (both editions):
    1. Title strip · Map strip · Quick Start
    2. Inputs reference (Full edition adds a CLASS Import workflow)
    3. Reading the Analyser + Comparison tearsheet
    4. Caveats + Troubleshooting
    5. Cheat sheet

Independent of v1 and v2 builders — neither imports the other.
"""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from div296 import __version__

REPO_ROOT = Path(__file__).resolve().parents[1]

# --- Colour palette ---------------------------------------------------------
INK = (0x1D, 0x3B, 0x34)        # deep sage — primary
SUB = (0x3F, 0x4F, 0x4A)        # slate — secondary
MUTE = (0x55, 0x55, 0x55)       # muted body
WARN = (0xA6, 0x1B, 0x1B)       # red — material caveats
AMBER = (0x8A, 0x6D, 0x00)      # amber — non-blocking warnings
COPPER = (0x8B, 0x4A, 0x1E)     # copper — Comparison section
TEAL = (0x2D, 0x5F, 0x5C)       # teal — CLASS Import (Full edition only)


def _hex(rgb: tuple[int, int, int]) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


# ============================================================================
# Paragraph + run helpers
# ============================================================================

def _run(p, text, *, size=10.5, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = RGBColor(*color)
    return r


def _para(doc, *, before=0, after=2, line_spacing=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    return p


def _page_break(doc):
    p = _para(doc, after=0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_margins(doc, *, top=1.2, bottom=1.2, left=1.4, right=1.4):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


# ============================================================================
# OxmlElement helpers
# ============================================================================

def _bookmark_id():
    _bookmark_id._n += 1
    return _bookmark_id._n


_bookmark_id._n = 100


def _add_bookmark(paragraph, name: str) -> None:
    bid = _bookmark_id()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _set_paragraph_left_border(paragraph, *, color, width_pt=4, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt * 8))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), _hex(color))
    p_bdr.append(left)
    p_pr.append(p_bdr)


def _shade_cell(cell, rgb):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(rgb))
    tc_pr.append(shd)


# ============================================================================
# Section / inline blocks
# ============================================================================

def h1(doc, text, *, accent=INK, bookmark_name: str | None = None):
    p = _para(doc, before=4, after=2)
    p.style = doc.styles["Heading 1"]
    _run(p, text, size=14, bold=True, color=accent)
    _set_paragraph_left_border(p, color=accent, width_pt=4, space=8)
    if bookmark_name:
        _add_bookmark(p, bookmark_name)
    return p


def h2(doc, text, *, color=SUB):
    p = _para(doc, before=4, after=1)
    p.style = doc.styles["Heading 2"]
    _run(p, text, size=11, bold=True, color=color)
    return p


def body(doc, text, *, size=10, italic=False, color=None):
    p = _para(doc, before=0, after=2, line_spacing=1.15)
    _run(p, text, size=size, italic=italic, color=color)
    return p


def bullets(doc, items, *, tight=True, size=10):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1 if tight else 2)
        if isinstance(item, tuple):
            head, tail = item
            _run(p, head, size=size, bold=True)
            _run(p, " — " + tail, size=size)
        else:
            _run(p, item, size=size)


def simple_table(doc, headers, rows, *, font_size=9, header_fill=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        _run(c.paragraphs[0], h, size=font_size, bold=True,
             color=(0xFF, 0xFF, 0xFF) if header_fill else None)
        if header_fill is not None:
            _shade_cell(c, header_fill)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ""
            _run(c.paragraphs[0], str(val), size=font_size)
    return t


# ============================================================================
# Page 1 — title strip + map strip + Quick Start
# ============================================================================

def title_strip(doc, *, include_class: bool) -> None:
    t = _para(doc, before=0, after=0)
    _run(t, "Division 296 Cost Base Reset Model", size=18, bold=True, color=INK)

    edition_label = "Full Edition" if include_class else "Adviser Edition"
    audience = (
        "For SMSF advisers using CLASS Super as administration software."
        if include_class else
        "For SMSF advisers who don't use CLASS Super."
    )

    sub = _para(doc, before=0, after=2)
    _run(sub, edition_label, size=11, bold=True, color=SUB)
    _run(sub, "   ·   ", size=10, color=MUTE)
    _run(sub, f"v{__version__}", size=10, color=MUTE)
    _run(sub, "   ·   ", size=10, color=MUTE)
    _run(sub, f"Built {date.today().isoformat()}", size=10, italic=True, color=MUTE)
    _run(sub, "   ·   ", size=10, color=MUTE)
    _run(sub, audience, size=10, italic=True, color=SUB)


def about_panel(doc, *, include_class: bool) -> None:
    """Four-line orientation block between title strip and map strip.

    Frames Division 296, the reset election, what the model compares, and the
    decision the tool supports. Designed so a brand-new reader has the 'why'
    before the 'how' on the very first page.
    """
    p = _para(doc, before=6, after=1)
    _run(p, "What this tool does", size=11, bold=True, color=INK)

    panel = doc.add_table(rows=4, cols=1)
    panel.autofit = True

    rows = [
        ("Division 296.",
         " From FY 2026–27 (first test 30 Jun 2027), super earnings on the "
         "slice of Total Super Balance (TSB) between $3m and $10m attract an "
         "extra 15% tax; the slice above $10m attracts an extra 25%. Assessed "
         "personally on the member, not on the fund."),
        ("Cost base reset election.",
         " A transitional one-off election: step each CGT asset's cost base "
         "up to its market value at 30 Jun 2026. Once lodged, irrevocable."),
        ("What this model compares.",
         " Reset election ON vs OFF. For each scenario it projects Division "
         "296 tax — and the ordinary CGT that flows from realising the asset."),
        ("Decision it supports.",
         " Whether the SMSF should lodge the reset election. Rule of thumb: "
         "reset usually saves tax on future-gain assets and costs tax on "
         "future-loss assets — net effect depends on the asset mix."),
    ]
    for i, (lead, tail) in enumerate(rows):
        cell = panel.rows[i].cells[0]
        cell.text = ""
        _shade_cell(cell, (0xF5, 0xF1, 0xEA))  # very light parchment
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(1)
        cp.paragraph_format.space_after = Pt(1)
        _run(cp, lead, size=9.5, bold=True, color=INK)
        _run(cp, tail, size=9.5, color=(0x33, 0x33, 0x33))


def map_strip(doc, *, include_class: bool) -> None:
    """Workbook map — 4 cards (adviser) or 5 cards (full, with CLASS Import)."""
    if include_class:
        tabs = [
            ("INPUTS",       "you edit",     INK),
            ("CLASS IMPORT", "paste here",   TEAL),
            ("ANALYSER",     "auto",         SUB),
            ("COMPARISON",   "auto + print", COPPER),
            ("NOTES",        "reference",    AMBER),
        ]
        # 5 tabs + 4 arrows = 9 cols
        n_cols = 9
        # Arrows between every adjacent pair EXCEPT the last (Notes is
        # reference, not in the flow). With CLASS the flow is:
        # INPUTS · CLASS IMPORT → ANALYSER → COMPARISON | NOTES
        # CLASS feeds INPUTS via copy-paste-values (not a live link), so
        # show INPUTS ↔ CLASS as a paired "data entry" group and arrows
        # from there into Analyser/Comparison.
        arrow_positions = (3, 5)   # between CLASS↔ANALYSER and ANALYSER↔COMPARISON
        empty_positions = (1, 7)   # between INPUTS-CLASS (paired) and COMPARISON-NOTES
    else:
        tabs = [
            ("INPUTS",     "you edit",       INK),
            ("ANALYSER",   "auto",           SUB),
            ("COMPARISON", "auto + print",   COPPER),
            ("NOTES",      "reference",      AMBER),
        ]
        n_cols = 7
        arrow_positions = (1, 3)
        empty_positions = (5,)

    t = doc.add_table(rows=1, cols=n_cols)
    t.autofit = True

    for i, (name, role, col) in enumerate(tabs):
        cell = t.rows[0].cells[i * 2]
        cell.text = ""
        _shade_cell(cell, col)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(0)
        # Slightly smaller font in 5-tab mode so labels fit
        font = 9 if include_class else 10
        _run(p1, name, size=font, bold=True, color=(0xFF, 0xFF, 0xFF))
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        _run(p2, role, size=7.5, italic=True, color=(0xFF, 0xFF, 0xFF))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in arrow_positions:
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "→", size=14, bold=True, color=MUTE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in empty_positions:
        t.rows[0].cells[i].text = ""

    if include_class:
        caption = _para(doc, before=2, after=0)
        _run(caption,
             "CLASS Import is a staging tab: paste a CLASS Investment Summary "
             "Report, then copy-paste-values into the Inputs asset register.",
             size=8, italic=True, color=MUTE)


def quick_start(doc, *, include_class: bool) -> None:
    h1(doc, "Quick Start", accent=INK, bookmark_name="qs_quickstart")
    body(doc, "Three steps. Five minutes. One headline number.",
         size=10.5, italic=True, color=MUTE)

    # --- Three steps as a styled table -----------------------------------
    step_table = doc.add_table(rows=3, cols=2)
    step_table.autofit = True

    if include_class:
        step2_head = "Asset register — type, or paste from CLASS Import."
        step2_body = (
            "For each CGT asset held at 30 June 2026: original cost base (C), "
            "market value at 30 Jun 2026 (E), projected sale proceeds (G), "
            "held > 12 months (I). Column H is locked. For CLASS Super users, "
            "the CLASS Import tab maps a pasted Investment Summary Report "
            "(Tax Cost Base basis) into A:G — copy + Paste-Special > Values."
        )
    else:
        step2_head = "Inputs → Asset register."
        step2_body = (
            "For each CGT asset held at 30 June 2026: original cost base (C), "
            "market value at 30 Jun 2026 (E), projected sale proceeds (G), "
            "held > 12 months (I). Column H is locked — don't edit."
        )

    step_data = [
        ("1", "Inputs → Members.",
         "Type each member's Total Super Balance (TSB) into column B "
         "(rows 7–10). Split %, $3m–$10m band, and above-$10m band "
         "derive automatically."),
        ("2", step2_head, step2_body),
        ("3", "Comparison tab.",
         "Read the three headline cards: If no reset · If elected · Net effect. "
         "Negative net effect (green) = the reset saves Division 296 tax."),
    ]
    for i, (num, head, tail) in enumerate(step_data):
        num_cell = step_table.rows[i].cells[0]
        text_cell = step_table.rows[i].cells[1]
        num_cell.width = Cm(1.0)
        num_cell.text = ""
        text_cell.text = ""
        _shade_cell(num_cell, INK)
        np = num_cell.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(np, num, size=16, bold=True, color=(0xFF, 0xFF, 0xFF))
        num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tp = text_cell.paragraphs[0]
        tp.paragraph_format.space_after = Pt(1)
        _run(tp, head, size=10.5, bold=True, color=INK)
        tp2 = text_cell.add_paragraph()
        tp2.paragraph_format.space_after = Pt(1)
        _run(tp2, tail, size=10)

    h2(doc, "Where the answer lives")
    body(doc,
         "Comparison tab → Headline cards. Three numbers: total Div 296 tax "
         "under each scenario plus the signed net effect. Negative (green) = "
         "reset saves tax. Positive (red) = ‘reset trap’ — locking in a higher "
         "basis on an asset you'll later sell at a loss shrinks that future "
         "loss, so net tax goes up.")

    h2(doc, "Two safety banners on Inputs")
    p = _para(doc, after=1)
    _run(p, "Red — formulas in the register. ", size=10, bold=True, color=WARN)
    formulas_tail = (
        "Pasted from CLASS Import with Ctrl+V instead of Paste-Special > Values. "
        "Press Ctrl+Z, re-paste as values."
        if include_class else
        "You used Ctrl+V. Press Ctrl+Z, re-paste as Paste-Special > Values."
    )
    _run(p, formulas_tail, size=10)
    p = _para(doc, after=1)
    _run(p, "Red — proceeds without 30 Jun 2026 value. ", size=10, bold=True, color=WARN)
    _run(p, "Rows with column G filled but column E blank are silently excluded. "
            "Fill column E, or clear column G.", size=10)

    _page_break(doc)


# ============================================================================
# Page 2 — Inputs reference (+ CLASS workflow in Full edition)
# ============================================================================

def inputs_reference(doc, *, include_class: bool) -> None:
    h1(doc, "Inputs reference", accent=INK, bookmark_name="qs_inputs")
    body(doc, "Column-level reference for the only tab whose numbers flow to "
              "tax calcs. Everything else in the workbook reads from here.",
         size=10, italic=True, color=MUTE)

    h2(doc, "Section 1 — Members (rows 6–12)")
    body(doc, "Up to four members. Edit columns A (label) and B (TSB, dollars). "
              "Three derived columns are read-only:")
    bullets(doc, [
        ("Split %", "TSB ÷ total fund TSB."),
        ("Band 1 (3m–10m)", "MAX(0, MIN(TSB, $10m) − $3m) ÷ TSB. Taxed at 15%."),
        ("Band 2 (> 10m)", "MAX(0, TSB − $10m) ÷ TSB. Taxed at 25%."),
    ])
    body(doc, "Traffic-light banner above the table reads green (no Div 296), "
              "amber ($3m–$10m), or deep amber (above $10m).")

    h2(doc, "Section 2 — Asset register (rows 15–65, 50 rows)")
    simple_table(doc, ["Col", "Header", "Notes"], [
        ("A", "Asset code", "Short identifier. Free text."),
        ("B", "Asset name", "Plain description."),
        ("C", "Original cost base",
         "Pre-reset basis — used for ordinary CGT regardless of reset."),
        ("D", "Current market value", "Context only; not used in tax calcs."),
        ("E", "Market value at 30 Jun 2026",
         "Load-bearing. Div 296 cost base under the reset scenario."),
        ("F", "Valuation source / date", "Free text; mirrored to Notes valuation log."),
        ("G", "Projected sale proceeds",
         "Blank = row excluded from totals (amber tripwire fires)."),
        ("H", "Projected gain/loss", "Locked formula. Do not edit."),
        ("I", "Held > 12 months?", "Yes/No dropdown. Drives 1/3 CGT discount (s115-100)."),
    ], header_fill=INK)

    h2(doc, "Section 3 — Advanced assumptions (rows 67–75)")
    body(doc, "Eight named-range cells. Touch only if the ATO publishes a change. "
              "Defaults match the enacted law:")
    bullets(doc, [
        "rate_tier1 = 15% (slice $3m–$10m). rate_tier2 = 25% (slice above $10m).",
        "threshold_1 = $3,000,000 · threshold_2 = $10,000,000.",
        "discount_rate = 1/3 (s115-100 ITAA 1997). fund_cgt_rate = 15% (accumulation).",
        "indexation_increment_1 / _2 — forward-modelling only, not applied to Year 1.",
    ])

    if include_class:
        h1(doc, "Using the CLASS Import tab", accent=TEAL,
           bookmark_name="qs_class_import")
        body(doc,
             "CLASS Import is a staging area that maps a CLASS Super Investment "
             "Summary Report into the shape of the Inputs asset register. It "
             "holds no live formulas into the model — data only reaches the "
             "register via a documented copy-paste-values step.",
             size=10)

        bullets(doc, [
            ("1. In CLASS",
             "Generate the Investment Summary Report on a TAX COST BASE basis "
             "(not Accounting — Accounting overstates basis for trusts/ETFs/managed "
             "funds and silently produces wrong tax). Export to CSV."),
            ("2. Clear the green PASTE ZONE first",
             "Select A7:R56 on the CLASS Import tab, press Delete. A red "
             "‘EXAMPLE DATA STILL PRESENT’ banner stays lit over the mapped "
             "block until every demo row is gone."),
            ("3. Paste the CLASS rows",
             "Paste your CLASS export rows into A7 onwards. The MAPPED BLOCK on "
             "the right filters out cash + the REASEDCGT realised-CGT line, then "
             "maps the survivors into register shape."),
            ("4. Copy mapped block, paste into Inputs",
             "Select the mapped block cols A:G (range T7:Z56), copy, then on "
             "Inputs paste into A16 using Paste-Special > Values (never plain "
             "Ctrl+V — the locked col H rejects an A:H paste)."),
            ("5. Complete by hand",
             "Fill Market value @ 30 Jun (E), Valuation source/date (F), "
             "Projected proceeds (G), Held > 12 months (I). Resolve any "
             "negative-cost-base flag (possible CGT event E4)."),
        ])

    _page_break(doc)


# ============================================================================
# Page 3 — Analyser + Comparison
# ============================================================================

def analyser_and_comparison(doc) -> None:
    h1(doc, "Reading the Analyser", accent=SUB, bookmark_name="qs_analyser")

    h2(doc, "Fund summary (rows 6–13)")
    body(doc, "Side by side: If no reset · If elected · Difference. Bottom row "
              "(‘Total Div 296 tax’) is the headline. Signed Difference: "
              "negative (green) = reset saves tax; positive (red) = reset costs more tax.")

    h2(doc, "Per-asset analysis (rows 15–67)")
    body(doc, "Always reflects the elected-reset scenario. Most-read columns:")
    bullets(doc, [
        ("K — Div 296 tax",
         "Authoritative per-asset Div 296 tax (pro-rata of the elected-reset headline)."),
        ("L — Reset impact",
         "Signed delta. Negative = reset reduces this asset's burden. "
         "Positive = reset-trap candidate."),
        ("E, G — Ord gross gain / Ord CGT (info only)",
         "Greyed. Do NOT sum to fund Ord CGT — s102-5 capital-loss netting happens "
         "at fund level. Read the Reconciliation panel for the authoritative figure."),
    ])

    h2(doc, "Reconciliation (rows 70–74)")
    body(doc, "Authoritative fund-level totals: ordinary CGT after intra-year "
              "capital-loss netting (s102-5 method, then s115-100 1/3 discount on "
              "the discountable residue), Division 296 tax payable, and any "
              "current-year net unused capital loss carried forward.")

    h1(doc, "Comparison tearsheet", accent=COPPER, bookmark_name="qs_comparison")
    body(doc, "Print-ready landscape A4. Use as a client meeting handout or file "
              "note without editing. Order top to bottom:")
    bullets(doc, [
        "Header strip — fund, prepared by, date, version.",
        "Members & TSB strip.",
        "Headline cards — Div 296 tax under each scenario plus signed net effect.",
        "Per-scenario subtotals — Earnings, Ord CGT (unchanged), Div 296 tax, "
        "Total burden.",
        "Per-member breakdown — TSB and Div 296 tax for both scenarios.",
        "Per-asset detail — top 10 by |Δ Div 296 tax|.",
    ])

    _page_break(doc)


# ============================================================================
# Page 4 — Caveats + Troubleshooting
# ============================================================================

def caveats_and_troubleshooting(doc, *, include_class: bool) -> None:
    h1(doc, "Caveats", accent=AMBER, bookmark_name="qs_caveats")
    body(doc, "Hold these in mind. Material; several work in opposing directions.",
         size=10, italic=True, color=MUTE)
    caveat_items = [
        ("Prior-year carry-forward capital losses NOT modelled.",
         "Apply prior-year carry-forward outside the model."),
        ("Pension phase NOT modelled.",
         "Assumes 100% accumulation. Overstates Ord CGT for pension members."),
        ("Reset-OFF scenario is realised-only.",
         "Under enacted Div 296, no-reset earnings are assessed on year-on-year "
         "TSB movement (realised + unrealised). Comparison understates the "
         "no-reset burden."),
        ("Multi-member splits are TSB-proportional.",
         "Real funds use an actuarial certificate on time-weighted balances."),
        ("Wash sale / Part IVA risk.",
         "Pre-30 June 2026 disposals purely for tax sit in anti-avoidance "
         "territory (TR 2008/1, Part IVA)."),
        ("Transaction costs not modelled.",
         "Brokerage, stamp duty, liquidity drag, market-timing excluded."),
        ("Sheet protection is tamper-evident, not tamper-proof.",
         "Passwordless. Prevents accidental overwrites only."),
    ]
    if include_class:
        caveat_items.insert(0, (
            "CLASS Import needs the TAX cost base export.",
            "An Accounting-basis export looks identical in CSV but silently "
            "overstates cost bases for trusts, ETFs, and managed funds — "
            "wrong tax. Negative tax cost bases are passed through and "
            "flagged (possible CGT event E4)."
        ))
    bullets(doc, caveat_items)

    h1(doc, "Troubleshooting", accent=WARN, bookmark_name="qs_troubleshooting")
    trouble_rows = [
        ("Yellow ‘Sample data preloaded’ badge persists.",
         "Sample input (P1/S1/L1 or seeded TSBs) still present.",
         "Overwrite every sample cell on Inputs."),
        ("Red banner: proceeds without 30 Jun 2026 value.",
         "Those rows excluded from every figure.",
         "Fill column E, or clear column G."),
        ("Red banner: formulas in the register.",
         "Non-values paste left live formulas in register cells.",
         "Ctrl+Z to undo, then Paste-Special > Values."),
        ("TSB diagnostic banner is deep amber.",
         "At least one member's TSB exceeds $10m — 25% above-$10m tier applies.",
         "Read Comparison for impact. No input change."),
        ("Per-asset Ord CGT (col G) doesn't sum to Reconciliation Ord CGT.",
         "Correct — loss netting at fund level (s102-5), not per asset.",
         "Rely on Reconciliation."),
    ]
    if include_class:
        trouble_rows.append((
            "CLASS Import: red EXAMPLE DATA STILL PRESENT banner.",
            "Demo rows weren't cleared from the paste zone before pasting.",
            "Select A7:R56 on CLASS Import, Delete, re-paste your CLASS export."
        ))
    simple_table(doc, ["You see", "What it means", "Do this"], trouble_rows,
                 header_fill=WARN)

    _page_break(doc)


# ============================================================================
# Page 5 — Cheat sheet
# ============================================================================

def cheat_sheet(doc, *, include_class: bool) -> None:
    h1(doc, "Cheat sheet", accent=INK, bookmark_name="qs_cheatsheet")
    body(doc, "Tear-off reference. Keep next to the workbook.",
         size=10, italic=True, color=MUTE)

    h2(doc, "Inputs columns at a glance")
    bullets(doc, [
        "A · Asset code (free text)",
        "B · Asset name (free text)",
        "C · Original cost base (ordinary CGT basis — unchanged by reset)",
        "D · Current market value (context only)",
        "E · MV @ 30 Jun 2026 (load-bearing — Div 296 basis under reset)",
        "F · Valuation source / date",
        "G · Projected sale proceeds (blank → row excluded)",
        "H · Projected gain/loss (locked formula — do not edit)",
        "I · Held > 12 months? (drives 1/3 CGT discount)",
    ], size=9.5)

    if include_class:
        h2(doc, "CLASS Import workflow (5 steps)")
        bullets(doc, [
            "Export Investment Summary Report from CLASS — TAX Cost Base basis.",
            "On CLASS Import tab: select A7:R56, Delete (clear demo data).",
            "Paste CLASS rows into A7 onwards.",
            "Copy mapped block T7:Z56 (cols A:G).",
            "Inputs!A16 → Paste-Special > Values. Fill E/F/G/I by hand.",
        ], size=9.5)

    h2(doc, "Banners — colour, meaning, action")
    banner_rows = [
        ("Yellow", "Sample data still loaded.", "Overwrite every sample cell."),
        ("Red", "Formulas in register (Ctrl+V used).", "Ctrl+Z, Paste-Special > Values."),
        ("Red", "Proceeds without 30 Jun 2026 value.", "Fill col E or clear col G."),
        ("Amber", "TSB > $10m — 25% tier applies.", "No action; read Comparison."),
    ]
    if include_class:
        banner_rows.append(
            ("Red", "CLASS Import: demo data still in paste zone.",
             "Delete A7:R56 on CLASS Import, re-paste.")
        )
    simple_table(doc, ["Colour", "Means", "Do this"], banner_rows)

    h2(doc, "Keyboard tricks")
    bullets(doc, [
        ("Paste-Special > Values", "Ctrl+Alt+V then V — only safe paste into register."),
        ("Undo", "Ctrl+Z — recovers from a bad paste."),
        ("Word Navigation Pane (this guide)", "View → Navigation Pane."),
    ], size=9.5)

    h2(doc, "Five caveats that change numbers materially")
    bullets(doc, [
        "Prior-year carry-forward losses — apply outside the model.",
        "Pension phase — not modelled (overstates Ord CGT for pension members).",
        "No-reset is realised-only — understates no-reset Div 296 burden.",
        "Multi-member split is TSB-proportional, not actuarial.",
        "Transaction costs (brokerage, stamp duty, liquidity) excluded.",
    ], size=9.5)

    h2(doc, "Verifying the version")
    body(doc, "Notes tab → unhide rows immediately below the valuation log to "
              "see build_version, build_date, git_short_sha.", size=9.5)


# ============================================================================
# Page header
# ============================================================================

def _stamp_header(doc, *, edition_label: str):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p, f"{edition_label}  ·  v{__version__}", size=8, italic=True, color=MUTE)


# ============================================================================
# Entry point
# ============================================================================

def build_guide(*, include_class: bool) -> Document:
    doc = Document()
    _set_margins(doc)
    _stamp_header(doc, edition_label="Full Edition" if include_class else "Adviser Edition")

    title_strip(doc, include_class=include_class)
    about_panel(doc, include_class=include_class)
    map_strip(doc, include_class=include_class)
    quick_start(doc, include_class=include_class)
    inputs_reference(doc, include_class=include_class)
    analyser_and_comparison(doc)
    caveats_and_troubleshooting(doc, include_class=include_class)
    cheat_sheet(doc, include_class=include_class)

    return doc


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Build the Div 296 user guide v3.")
    parser.add_argument(
        "--edition",
        choices=("full", "adviser"),
        default="adviser",
        help="'adviser' (default) omits CLASS Import content; "
             "'full' includes CLASS Import workflow + caveats + troubleshooting.",
    )
    parser.add_argument(
        "--output", "-o",
        type=Path,
        default=None,
        help="Output path (default: docs/USER_GUIDE_{Adviser|Full}_Edition_v3.docx).",
    )
    args = parser.parse_args(argv)

    include_class = args.edition == "full"
    label = "Full" if include_class else "Adviser"
    out_path = args.output or (REPO_ROOT / "docs" / f"USER_GUIDE_{label}_Edition_v3.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = build_guide(include_class=include_class)
    doc.save(out_path)
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
