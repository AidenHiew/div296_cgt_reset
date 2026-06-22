"""Build the Adviser Edition user guide (.docx).

    python scripts/build_user_guide.py
    # writes docs/USER_GUIDE_Adviser_Edition.docx

Structure:
    Page 1   Quick-start tearsheet — everything an adviser needs to
             kickstart a model run without reading further.
    Page 2+  Depth — Inputs reference, reading the Analyser, the
             Comparison tearsheet, caveats, troubleshooting.

Voice: master copywriter (lead with the answer, imperative,
short sentences) and tax accountant (precise statutory references,
no soft-pedalling on caveats).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt, RGBColor

from div296 import __version__

REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs" / "USER_GUIDE_Adviser_Edition.docx"

INK = (0x1D, 0x3B, 0x34)        # primary heading colour (deep sage)
SUB = (0x3F, 0x4F, 0x4A)        # sub-heading colour
MUTE = (0x55, 0x55, 0x55)       # muted body
WARN = (0xA6, 0x1B, 0x1B)       # red — material caveats
AMBER = (0x8A, 0x6D, 0x00)      # amber — non-blocking warnings


# --- low-level helpers ------------------------------------------------------

def _run(p, text, *, size=10.5, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = RGBColor(*color)
    return r


def _para(doc, *, before=0, after=2, line_spacing=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    return p


def _set_margins(doc, *, top=1.5, bottom=1.5, left=1.7, right=1.7):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def _page_break(doc):
    p = _para(doc, after=0)
    p.add_run().add_break(WD_BREAK.PAGE)


# --- structural blocks ------------------------------------------------------

def title_block(doc, title, subtitle):
    p = _para(doc, before=0, after=2)
    _run(p, title, size=20, bold=True, color=INK)
    p2 = _para(doc, before=0, after=4)
    _run(p2, subtitle, size=10.5, italic=True, color=MUTE)


def lead(doc, text):
    """One-sentence value-prop, prominent."""
    p = _para(doc, before=2, after=6, line_spacing=1.15)
    _run(p, text, size=11.5, bold=True, color=INK)


def banner(doc, text, *, color=AMBER):
    p = _para(doc, before=2, after=4)
    p.paragraph_format.left_indent = Cm(0.3)
    _run(p, text, size=10, italic=True, color=color)


def h1(doc, text):
    p = _para(doc, before=10, after=2)
    _run(p, text, size=13, bold=True, color=INK)


def h2(doc, text):
    p = _para(doc, before=6, after=1)
    _run(p, text, size=11, bold=True, color=SUB)


def body(doc, text):
    p = _para(doc, before=0, after=3, line_spacing=1.15)
    _run(p, text, size=10.5)


def bullets(doc, items, *, tight=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1 if tight else 3)
        if isinstance(item, tuple):
            head, tail = item
            _run(p, head, size=10.5, bold=True)
            _run(p, " — " + tail, size=10.5)
        else:
            _run(p, item, size=10.5)


def steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            head, tail = item
            _run(p, head, size=10.5, bold=True)
            _run(p, " — " + tail, size=10.5)
        else:
            _run(p, item, size=10.5)


def table(doc, headers, rows, *, font_size=9.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        _run(c.paragraphs[0], h, size=font_size, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ""
            _run(c.paragraphs[0], str(val), size=font_size)
    return t


# --- page 1: quick-start tearsheet -----------------------------------------

def page_one(doc):
    title_block(
        doc,
        "Division 296 Cost Base Reset Model",
        f"Adviser Edition · v{__version__} · Quick Start",
    )

    lead(doc,
         "Model the impact of the Division 296 cost-base reset election on "
         "an SMSF in five minutes. One inputs tab, two scenarios side by side, "
         "one printable client tearsheet.")

    banner(doc,
           "⚠ Sample data is preloaded. Overwrite every yellow cell with the "
           "fund's actual figures before sharing with a client.")

    h1(doc, "Three steps to a result")
    steps(doc, [
        ("Inputs → Members (rows 7–10).",
         "Type each member's TSB into column B. Split %, $3m–$10m band, and "
         "above-$10m band derive automatically."),
        ("Inputs → Asset register (rows 16–65).",
         "For each CGT asset held at 30 June 2026: original cost base (C), "
         "market value at 30 Jun 2026 (E), projected sale proceeds (G), "
         "held > 12 months (I). Column H (Projected gain/loss) is locked — "
         "it's a formula."),
        ("Comparison tab.",
         "Read the three headline cards: If no reset · If elected · Net effect. "
         "Negative net effect (green) = the reset saves Division 296 tax."),
    ])

    h1(doc, "The four tabs")
    bullets(doc, [
        ("Inputs", "the only tab you edit."),
        ("Analyser", "fund summary, per-member tax, per-asset detail, "
         "reconciliation. Read-only."),
        ("Comparison", "print-ready one-page tearsheet for the client meeting."),
        ("Notes", "terminology, full caveats, valuation log. Royal Assent "
         "date cell is editable."),
    ], tight=True)

    h1(doc, "Where the answer lives")
    body(doc,
         "Comparison tab → Headline cards (top of the tearsheet). Three numbers: "
         "total Div 296 tax under each scenario plus the signed net effect.")
    bullets(doc, [
        ("Negative net effect", "reset reduces Division 296 tax — typically "
         "on assets sitting at an unrealised gain at 30 June 2026."),
        ("Positive net effect", "reset increases Division 296 tax — the "
         "‘reset trap’ on assets at an unrealised loss. Read these per-asset "
         "in the Analyser tab, column L."),
    ], tight=True)

    h1(doc, "Two safety banners to watch on Inputs")
    bullets(doc, [
        ("Red — formulas in the register.",
         "You used Ctrl+V instead of Paste-Special > Values. Press Ctrl+Z, "
         "re-paste as values."),
        ("Red — proceeds without 30 Jun 2026 value.",
         "Any row with column G filled but column E blank is silently excluded "
         "from every figure. Fill column E, or clear column G."),
    ], tight=True)

    banner(doc,
           "Illustrative only — not financial, tax or legal advice. Confirm "
           "every figure against the final ATO method and the fund's tax "
           "position before relying on it.",
           color=WARN)

    _page_break(doc)


# --- page 2+: depth ---------------------------------------------------------

def detail_section(doc):
    title_block(
        doc,
        "Detail & Reference",
        "Read this if you want more than the quick-start — column-level "
        "Inputs reference, how the Analyser is structured, the caveats that "
        "shape every number, and a troubleshooting table.",
    )

    # --- 1. What the model does ------------------------------------------
    h1(doc, "1. What the model does")
    body(doc,
        "The workbook compares two mutually exclusive scenarios at the "
        "30 June 2026 reset election date:")
    bullets(doc, [
        ("If no reset (default)",
         "every CGT asset keeps its original cost base for both ordinary "
         "CGT and Division 296 calculations."),
        ("If elected to reset",
         "every CGT asset's Division 296 cost base is stepped to its market "
         "value at 30 June 2026. Ordinary CGT cost base is unaffected."),
    ])
    body(doc,
        "The election is one-off, all-or-nothing, and irrevocable — it applies "
        "to every CGT asset held by the fund at 30 June 2026, or none of "
        "them. The model never recommends an outcome; it presents the signed "
        "difference and lets the adviser make the call.")

    # --- 2. Inputs reference --------------------------------------------
    h1(doc, "2. Inputs reference")

    h2(doc, "Section 1 — Members (rows 6–12)")
    body(doc,
         "Up to four members. Edit columns A (label, free text) and B (TSB, "
         "dollars). The three derived columns are read-only:")
    bullets(doc, [
        ("Split %", "TSB ÷ total fund TSB."),
        ("Band 1 (3m–10m)", "MAX(0, MIN(TSB, $10m) − $3m) ÷ TSB. Taxed at 15%."),
        ("Band 2 (> 10m)", "MAX(0, TSB − $10m) ÷ TSB. Taxed at 25%."),
    ], tight=True)
    body(doc,
         "Traffic-light banner above the Members table reads green (no Div 296), "
         "amber ($3m–$10m), or deep amber (above $10m). Use it as the "
         "applicability check before reading any other number.")

    h2(doc, "Section 2 — Asset register (rows 15–65, 50 rows)")
    table(doc, ["Col", "Header", "Notes"], [
        ("A", "Asset code",
         "Short identifier. Free text."),
        ("B", "Asset name",
         "Plain description (e.g. ‘Commercial property’)."),
        ("C", "Original cost base",
         "Pre-reset cost base — used for ordinary CGT regardless of the "
         "reset election."),
        ("D", "Current market value",
         "Context only; not used in tax calcs. Useful sanity-check for reviewers."),
        ("E", "Market value at 30 Jun 2026",
         "Load-bearing. Becomes the Div 296 cost base under the reset scenario. "
         "Garbage in, garbage out."),
        ("F", "Valuation source / date",
         "Free text. Mirrored verbatim into the Notes valuation log for audit trail."),
        ("G", "Projected sale proceeds",
         "Expected cash on disposal. Blank = row excluded from totals (and "
         "the amber tripwire fires)."),
        ("H", "Projected gain/loss",
         "Locked formula = proceeds − original cost base. Do not edit."),
        ("I", "Held > 12 months?",
         "Yes / No dropdown. Drives the 1/3 CGT discount under s115-100 ITAA 1997."),
    ])

    h2(doc, "Section 3 — Advanced assumptions (rows 67–75)")
    body(doc,
         "Eight named-range cells that feed every downstream formula. Touch only "
         "if the ATO publishes a change. Defaults match the enacted law at the "
         "build date:")
    bullets(doc, [
        "rate_tier1 = 15% (slice $3m–$10m).",
        "rate_tier2 = 25% (slice above $10m).",
        "threshold_1 = $3,000,000 · threshold_2 = $10,000,000.",
        "discount_rate = 1/3 (s115-100 ITAA 1997).",
        "fund_cgt_rate = 15% (accumulation phase).",
        "indexation_increment_1 / _2 (forward-modelling only — not applied to Year 1).",
    ], tight=True)

    # --- 3. Reading the Analyser ----------------------------------------
    h1(doc, "3. Reading the Analyser")

    h2(doc, "Fund summary (rows 6–13)")
    body(doc,
        "Side by side: If no reset · If elected · Difference. Bottom row "
        "(‘Total Div 296 tax’) is the headline. The Difference column is "
        "signed; negative (green) = reset saves tax, positive (red) = reset "
        "costs more tax.")

    h2(doc, "Per-asset analysis (rows 15–67)")
    body(doc,
        "Always reflects the elected-reset scenario. Most-read columns:")
    bullets(doc, [
        ("K — Div 296 tax",
         "Authoritative per-asset Div 296 tax (pro-rata of the elected-reset "
         "headline)."),
        ("L — Reset impact",
         "Signed delta. Negative = reset reduces this asset's Div 296 burden. "
         "Positive = reset-trap candidate."),
        ("E, G — Ord gross gain / Ord CGT (info only)",
         "Greyed. Do NOT sum to the fund Ord CGT because s102-5 capital-loss "
         "netting happens at the fund level, not per asset. Read the "
         "Reconciliation panel for the authoritative number."),
    ], tight=True)

    h2(doc, "Reconciliation (rows 70–74)")
    body(doc,
         "Authoritative fund-level totals: ordinary CGT after intra-year "
         "capital-loss netting (s102-5 method, then s115-100 1/3 discount on "
         "the discountable residue), Division 296 tax payable, and any current-"
         "year net unused capital loss carried forward.")

    # --- 4. Comparison tab ----------------------------------------------
    h1(doc, "4. Comparison tab — the client tearsheet")
    body(doc,
        "Print-ready landscape A4. Suitable as a client meeting handout or "
        "file note without editing. Order top to bottom:")
    bullets(doc, [
        "Header strip — fund, prepared by, date, version.",
        "Members & TSB strip.",
        "Headline cards — Div 296 tax under each scenario plus signed net effect.",
        "Per-scenario subtotals — Earnings, Ord CGT (unchanged), Div 296 tax, "
        "Total burden.",
        "Per-member breakdown — TSB and Div 296 tax for both scenarios.",
        "Per-asset detail — top 10 by |Δ Div 296 tax|. Overflow note points "
        "at the Analyser for the full register.",
    ], tight=True)

    # --- 5. Caveats that shape every number -----------------------------
    h1(doc, "5. Caveats that shape every number")
    body(doc,
         "Hold these in mind whenever you read a figure. They are material "
         "and several work in opposing directions.")
    bullets(doc, [
        ("Prior-year carry-forward capital losses are NOT modelled.",
         "The workbook nets gains and losses within the current income year "
         "(s102-5) but takes no brought-forward balance as input. Apply any "
         "prior-year carry-forward outside the model before relying on the "
         "ordinary CGT figure."),
        ("Pension phase is NOT modelled.",
         "The model assumes 100% accumulation phase (fund earnings tax = 15%). "
         "For funds with pension members it overstates ordinary CGT and the "
         "reset's relative attractiveness."),
        ("Reset-OFF scenario is realised-only.",
         "Under enacted Division 296, fund earnings without the reset are "
         "assessed on year-on-year TSB movement (realised + unrealised). The "
         "Comparison tab compares realised vs realised, so it understates the "
         "no-reset Division 296 burden."),
        ("Multi-member splits are TSB-proportional.",
         "A real fund determines each member's share via an actuarial "
         "certificate on time-weighted average balances. The model "
         "approximates with current-TSB share."),
        ("Wash sale / Part IVA risk.",
         "Pre-30 June 2026 disposals undertaken purely for the tax outcome "
         "sit in anti-avoidance territory (TR 2008/1, Part IVA). Document the "
         "commercial rationale separately."),
        ("Transaction costs not modelled.",
         "Brokerage, stamp duty (for property), liquidity drag, and "
         "market-timing risk are excluded."),
        ("Alternative levers not modelled.",
         "TSB recontribution / withdrawal splits, pension commencement, and "
         "estate-timing considerations can materially change the optimal "
         "outcome. None are reflected."),
        ("Sheet protection is tamper-evident, not tamper-proof.",
         "Protection (where applied) is passwordless. It prevents accidental "
         "overwrites; it does not enforce immutability."),
    ])

    # --- 6. Troubleshooting --------------------------------------------
    h1(doc, "6. Troubleshooting")
    table(doc, ["You see", "What it means", "Do this"], [
        ("Yellow ‘Sample data preloaded’ badge persists.",
         "At least one sample input (P1/S1/L1 codes or the seeded member "
         "TSBs) still present.",
         "Overwrite every sample cell on Inputs; the badge clears automatically."),
        ("Red banner: ‘Some rows have Projected sale proceeds but no Market "
         "value at 30 Jun 2026’.",
         "Those rows are excluded from every figure.",
         "Fill column E for the row, or clear column G."),
        ("Red banner: ‘Formulas detected in the asset register’.",
         "A non-values paste left live formulas in the register cells.",
         "Press Ctrl+Z to undo, then Paste-Special > Values."),
        ("TSB diagnostic banner is deep amber.",
         "At least one member's TSB exceeds $10m — the 25% above-$10m tier "
         "applies in addition to the 15% $3m–$10m tier.",
         "Read the Comparison tearsheet for the modelled impact — no input change needed."),
        ("Per-asset Ord CGT (col G) doesn't sum to Reconciliation Ord CGT.",
         "Correct behaviour. Capital-loss netting happens at the fund level "
         "(s102-5), not per asset.",
         "Rely on the Reconciliation panel for the authoritative figure."),
    ])

    # --- 7. About this guide -------------------------------------------
    h1(doc, "7. About this guide")
    body(doc,
        f"Built for Division 296 Cost Base Reset Model v{__version__} "
        "(Adviser Edition). The workbook carries a hidden provenance block "
        "on the Notes tab (build version, build date, source commit). Unhide "
        "the rows immediately below the valuation log on Notes to verify a "
        "received file against this guide.")
    banner(doc,
           "Illustrative only — not financial, tax or legal advice. Confirm "
           "every figure against the final ATO method and the fund's tax "
           "position before relying on it.",
           color=WARN)


# --- entry point -----------------------------------------------------------

def build_guide() -> Document:
    doc = Document()
    _set_margins(doc, top=1.5, bottom=1.5, left=1.7, right=1.7)
    page_one(doc)
    detail_section(doc)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_guide()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
