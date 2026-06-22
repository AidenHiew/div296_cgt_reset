"""Build the Adviser Edition user guide v2 (.docx).

    python scripts/build_user_guide_v2.py
    # writes docs/USER_GUIDE_Adviser_Edition_v2.docx

v2 over v1:
- Uses Word's built-in Heading 1/2 styles → Navigation Pane works automatically.
- Auto-TOC field on page 4 + manual hyperlinked fallback list.
- Bookmarks at every H1; cross-references hyperlink to them.
- Per-section left-edge accent band — geometry first, colour second
  (monochrome print still navigable).
- Cover poster + workbook visual map + cheat sheet back-cover.

Independent of v1's builder — neither imports the other.

Design notes:
- Section breaks must be CONTINUOUS, or every H1 forces a new page
  and the document balloons.
- Word's TOC and PAGEREF fields need a field-update on first open;
  the manual fallback list immediately below the TOC field guarantees
  navigation works either way.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from div296 import __version__

REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs" / "USER_GUIDE_Adviser_Edition_v2.docx"

# --- Colour palette (RGB tuples; copper is new to v2) -----------------------
INK = (0x1D, 0x3B, 0x34)        # deep sage — primary
SUB = (0x3F, 0x4F, 0x4A)        # slate — secondary
MUTE = (0x55, 0x55, 0x55)       # muted body
WARN = (0xA6, 0x1B, 0x1B)       # red — material caveats
AMBER = (0x8A, 0x6D, 0x00)      # amber — non-blocking warnings
COPPER = (0x8B, 0x4A, 0x1E)     # copper — Comparison section (v2 only)

ACCENT_NEAR_WHITE = (0xF6, 0xF6, 0xF4)


def _hex(rgb: tuple[int, int, int]) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


# ============================================================================
# Low-level paragraph + run helpers
# ============================================================================

def _run(p, text, *, size=10.5, bold=False, italic=False, color=None,
         underline=False):
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    if color is not None:
        r.font.color.rgb = RGBColor(*color)
    return r


def _para(doc, *, before=0, after=2, line_spacing=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    return p


def _page_break(doc):
    p = _para(doc, after=0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_margins(doc, *, top=1.4, bottom=1.4, left=1.6, right=1.6):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


# ============================================================================
# OxmlElement helpers — the stuff python-docx doesn't expose natively
# ============================================================================

def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_hyperlink(paragraph, anchor: str, text: str, *, size=10.5,
                   color=INK, bold=False):
    """Internal hyperlink (jumps to a bookmark)."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    rcolor = OxmlElement("w:color")
    rcolor.set(qn("w:val"), _hex(color))
    r_pr.append(rcolor)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    r_pr.append(rstyle)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_pageref(paragraph, bookmark: str, *, size=10.5, color=MUTE):
    """Insert a PAGEREF field — Word resolves it to the page number of the
    given bookmark on field-update. Until first update Word shows the
    field-result placeholder (e.g. '1'); the manual Contents page on p.4
    hard-codes page numbers in visible text as the always-correct fallback."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' PAGEREF {bookmark} \\h '
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    placeholder = OxmlElement("w:t")
    placeholder.text = "?"
    run._r.append(placeholder)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)


def _insert_toc_field(paragraph):
    """Insert a TOC field — Word populates on field-update."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click → Update Field to populate this Contents."
    run._r.append(placeholder)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*MUTE)


def _set_paragraph_left_border(paragraph, *, color, width_pt=4, space=8):
    """Coloured left rule on the paragraph — the section accent band."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt * 8))   # eighths of a point
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), _hex(color))
    p_bdr.append(left)
    p_pr.append(p_bdr)


def _shade_cell(cell, rgb):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(rgb))
    tc_pr.append(shd)


def _box_paragraph(paragraph, *, color):
    """Thin coloured box around a paragraph — used for the disclaimer stamp."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "8")
        bdr.set(qn("w:space"), "6")
        bdr.set(qn("w:color"), _hex(color))
        p_bdr.append(bdr)
    p_pr.append(p_bdr)


# ============================================================================
# Higher-level building blocks
# ============================================================================

def _next_bookmark_id():
    """Generator-style: every call yields a fresh integer."""
    _next_bookmark_id._n += 1
    return _next_bookmark_id._n


_next_bookmark_id._n = 100   # initialise — keeps IDs well away from 0


def h1(doc, text, *, accent=INK, bookmark_name: str | None = None):
    """A real Word Heading 1 — feeds the Nav Pane and the TOC field.

    Override font/colour at the RUN level (not the style level), so the
    paragraph keeps its Heading 1 identity for nav/TOC purposes.
    """
    p = _para(doc, before=6, after=3)
    p.style = doc.styles["Heading 1"]
    _run(p, text, size=16, bold=True, color=accent)
    _set_paragraph_left_border(p, color=accent, width_pt=4, space=8)
    if bookmark_name:
        _add_bookmark(p, bookmark_name, _next_bookmark_id())
    return p


def h2(doc, text, *, color=SUB):
    p = _para(doc, before=8, after=2)
    p.style = doc.styles["Heading 2"]
    _run(p, text, size=12, bold=True, color=color)
    return p


def body(doc, text, *, size=10.5, italic=False, color=None):
    p = _para(doc, before=0, after=3, line_spacing=1.18)
    _run(p, text, size=size, italic=italic, color=color)
    return p


def banner(doc, text, *, color=AMBER):
    p = _para(doc, before=3, after=4)
    p.paragraph_format.left_indent = Cm(0.3)
    _run(p, text, size=10, italic=True, color=color)
    return p


def bullets(doc, items, *, tight=False, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1 if tight else 3)
        if isinstance(item, tuple):
            head, tail = item
            _run(p, head, size=size, bold=True)
            _run(p, " — " + tail, size=size)
        else:
            _run(p, item, size=size)


def simple_table(doc, headers, rows, *, font_size=9.5, header_fill=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        _run(c.paragraphs[0], h, size=font_size, bold=True)
        if header_fill is not None:
            _shade_cell(c, header_fill)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ""
            _run(c.paragraphs[0], str(val), size=font_size)
    return t


# ============================================================================
# Page 1 — Cover
# ============================================================================

def cover(doc) -> None:
    # Push title down a little for poster feel
    spacer = _para(doc, before=0, after=0)
    spacer.paragraph_format.space_after = Pt(28)

    title = _para(doc, before=0, after=6)
    _run(title, "Division 296", size=30, bold=True, color=INK)

    title2 = _para(doc, before=0, after=14)
    _run(title2, "Cost Base Reset Model", size=26, bold=True, color=INK)

    sub = _para(doc, before=0, after=18)
    _run(sub, "Adviser Edition", size=14, bold=True, color=SUB)
    _run(sub, "   ·   ", size=14, color=MUTE)
    _run(sub, f"v{__version__}", size=14, color=MUTE)
    _run(sub, "   ·   ", size=14, color=MUTE)
    _run(sub, f"Built {date.today().isoformat()}", size=14, italic=True, color=MUTE)

    aud = _para(doc, before=0, after=12)
    _run(aud, "For SMSF advisers who don't use CLASS Super as administration software.",
         size=11, italic=True, color=SUB)

    lead = _para(doc, before=0, after=24)
    _run(lead, "Model the impact of the Division 296 cost-base reset election on an "
               "SMSF in five minutes. One inputs tab, two scenarios side by side, one "
               "printable client tearsheet.",
         size=12, color=INK)

    # Quick Contents strip
    qc = _para(doc, before=0, after=10)
    _run(qc, "QUICK CONTENTS  ", size=9, bold=True, color=MUTE)
    parts = [
        ("Quick Start", "qs_quickstart", "p.2"),
        ("Workbook map", "qs_map", "p.3"),
        ("Inputs reference", "qs_inputs", "p.5"),
        ("Cheat sheet", "qs_cheatsheet", "p.9"),
    ]
    for i, (label, anchor, page) in enumerate(parts):
        if i > 0:
            _run(qc, "   ·   ", size=10, color=MUTE)
        _add_hyperlink(qc, anchor, label, size=10, color=INK, bold=True)
        _run(qc, f" {page}", size=10, color=MUTE)

    # Push the disclaimer to the foot of the page
    pusher = _para(doc, before=0, after=0)
    pusher.paragraph_format.space_after = Pt(220)

    stamp = _para(doc, before=0, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(stamp, "ILLUSTRATIVE ONLY — NOT FINANCIAL, TAX OR LEGAL ADVICE.   "
                "Confirm every figure against the final ATO method and the fund's "
                "tax position before relying on it.",
         size=9, bold=True, color=WARN)
    _box_paragraph(stamp, color=WARN)

    _page_break(doc)


# ============================================================================
# Page 2 — Quick Start
# ============================================================================

def quick_start(doc) -> None:
    h1(doc, "Quick Start", accent=INK, bookmark_name="qs_quickstart")
    body(doc, "Three steps. Five minutes. One headline number.",
         size=11, italic=True, color=MUTE)

    # --- Three steps as a styled table -----------------------------------
    h2(doc, "Three steps to a result")
    step_table = doc.add_table(rows=3, cols=2)
    step_table.autofit = True
    step_data = [
        ("1", "Inputs → Members (rows 7–10).",
         "Type each member's TSB into column B. Split %, $3m–$10m band, "
         "and above-$10m band derive automatically."),
        ("2", "Inputs → Asset register (rows 16–65).",
         "For each CGT asset held at 30 June 2026: original cost base (C), "
         "market value at 30 Jun 2026 (E), projected sale proceeds (G), "
         "held > 12 months (I). Column H is a locked formula — don't edit."),
        ("3", "Comparison tab.",
         "Read the three headline cards: If no reset · If elected · Net effect. "
         "Negative net effect (green) = the reset saves Division 296 tax."),
    ]
    for i, (num, head, tail) in enumerate(step_data):
        num_cell = step_table.rows[i].cells[0]
        text_cell = step_table.rows[i].cells[1]
        num_cell.width = Cm(1.2)
        num_cell.text = ""
        text_cell.text = ""
        _shade_cell(num_cell, INK)
        np = num_cell.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(np, num, size=18, bold=True, color=(0xFF, 0xFF, 0xFF))
        num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tp = text_cell.paragraphs[0]
        tp.paragraph_format.space_after = Pt(2)
        _run(tp, head, size=11, bold=True, color=INK)
        tp2 = text_cell.add_paragraph()
        tp2.paragraph_format.space_after = Pt(2)
        _run(tp2, tail, size=10.5)

    # --- 4-tab grid (2x2) -------------------------------------------------
    h2(doc, "The four tabs")
    tab_table = doc.add_table(rows=2, cols=2)
    tab_table.autofit = True
    tab_data = [
        [("Inputs", "the only tab you edit.", INK),
         ("Analyser", "auto. Fund summary, per-asset detail, reconciliation.", SUB)],
        [("Comparison", "auto + printable client tearsheet.", COPPER),
         ("Notes", "terminology, full caveats, valuation log.", AMBER)],
    ]
    for r in range(2):
        for c in range(2):
            cell = tab_table.rows[r].cells[c]
            cell.text = ""
            name, desc, col = tab_data[r][c]
            p1 = cell.paragraphs[0]
            p1.paragraph_format.space_after = Pt(2)
            _run(p1, name, size=12, bold=True, color=col)
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            _run(p2, desc, size=10)

    # --- Where the answer lives ------------------------------------------
    h2(doc, "Where the answer lives")
    body(doc, "Comparison tab → Headline cards (top of the tearsheet). Three "
              "numbers: total Div 296 tax under each scenario plus the signed net effect.")
    bullets(doc, [
        ("Negative net effect", "reset reduces Division 296 tax — typically "
         "on assets sitting at an unrealised gain at 30 June 2026."),
        ("Positive net effect", "reset increases Division 296 tax — the "
         "‘reset trap’ on assets at an unrealised loss."),
    ], tight=True)

    # --- Safety banners ---------------------------------------------------
    h2(doc, "Two safety banners to watch on Inputs")
    p = _para(doc, after=2)
    _run(p, "Red — formulas in the register. ", bold=True, color=WARN)
    _run(p, "You used Ctrl+V instead of Paste-Special > Values. Press Ctrl+Z, "
            "re-paste as values.")
    p = _para(doc, after=2)
    _run(p, "Red — proceeds without 30 Jun 2026 value. ", bold=True, color=WARN)
    _run(p, "Any row with column G filled but column E blank is silently "
            "excluded from every figure. Fill column E, or clear column G.")

    _page_break(doc)


# ============================================================================
# Page 3 — Workbook at a glance
# ============================================================================

def workbook_map(doc) -> None:
    h1(doc, "Workbook at a glance", accent=INK, bookmark_name="qs_map")
    body(doc, "The four tabs, what each does, and how data flows between them.",
         size=11, italic=True, color=MUTE)

    # 7 columns: 4 tab cells + 3 arrow cells between them
    t = doc.add_table(rows=3, cols=7)
    t.autofit = True

    tabs = [
        ("INPUTS", "you edit",
         ["Members", "Asset register", "Assumptions"], INK),
        ("ANALYSER", "auto",
         ["Fund summary", "Per-asset detail", "Reconciliation"], SUB),
        ("COMPARISON", "auto + print",
         ["Headline cards", "Per-member breakdown", "Top-10 assets"], COPPER),
        ("NOTES", "reference",
         ["Terminology", "Caveats", "Valuation log"], AMBER),
    ]

    # Row 0: tab names
    for i, (name, _role, _contents, col) in enumerate(tabs):
        cell = t.rows[0].cells[i * 2]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(cell, col)
        _run(p, name, size=12, bold=True, color=(0xFF, 0xFF, 0xFF))
    # Arrow cells (only between first 3 tabs, last gap stays empty)
    for i in (1, 3):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "→", size=18, bold=True, color=MUTE)
    # Last arrow cell stays empty (Notes is reference, not in the flow)
    t.rows[0].cells[5].text = ""

    # Row 1: role caption
    for i, (_name, role, _contents, col) in enumerate(tabs):
        cell = t.rows[1].cells[i * 2]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(cell, ACCENT_NEAR_WHITE)
        _run(p, role, size=9, italic=True, color=col)
    for i in (1, 3, 5):
        t.rows[1].cells[i].text = ""

    # Row 2: contents bullets
    for i, (_name, _role, contents, col) in enumerate(tabs):
        cell = t.rows[2].cells[i * 2]
        cell.text = ""
        for k, item in enumerate(contents):
            p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _run(p, "• " + item, size=10)
    for i in (1, 3, 5):
        t.rows[2].cells[i].text = ""

    # Caption underneath
    cap = _para(doc, before=14, after=0)
    _run(cap, "Read the arrows: Inputs feeds the Analyser; the Analyser drives "
              "the Comparison tearsheet. Notes sits to the side as reference — "
              "it never feeds a tax figure.",
         size=10, italic=True, color=MUTE)

    _page_break(doc)


# ============================================================================
# Page 4 — Contents
# ============================================================================

# (anchor, label, page)  — manual fallback table; page numbers are visible
# (so they work in print) and the entries are clickable bookmarks.
TOC_ENTRIES = [
    ("qs_quickstart", "Quick Start", 2),
    ("qs_map", "Workbook at a glance", 3),
    ("qs_inputs", "Inputs reference", 5),
    ("qs_analyser", "Reading the Analyser", 7),
    ("qs_comparison", "Comparison tearsheet", 7),
    ("qs_caveats", "Caveats", 8),
    ("qs_troubleshooting", "Troubleshooting", 8),
    ("qs_cheatsheet", "Cheat sheet", 9),
]


def contents_page(doc) -> None:
    h1(doc, "Contents", accent=SUB, bookmark_name="qs_contents")
    body(doc, "Click any entry to jump to that section. The auto-TOC below "
              "populates when Word updates fields (right-click → Update Field).",
         size=10, italic=True, color=MUTE)

    # Auto-TOC field
    h2(doc, "Auto Table of Contents")
    p = _para(doc, after=10)
    _insert_toc_field(p)

    # Manual fallback list — always works, even in headless renders
    h2(doc, "Manual fallback (always works)")
    body(doc, "Use this if the auto-TOC above shows the placeholder text "
              "instead of a list.", size=10, color=MUTE)

    for anchor, label, page in TOC_ENTRIES:
        p = _para(doc, after=3, line_spacing=1.2)
        _add_hyperlink(p, anchor, label, size=11, color=INK, bold=True)
        # Dotted leader using spaces — Word's "tab-leader" is more work and
        # the simple visible page number reads just as well.
        _run(p, f"     …     p.{page}", size=11, color=MUTE)

    _page_break(doc)


# ============================================================================
# Page 5–6 — Inputs reference
# ============================================================================

def inputs_reference(doc) -> None:
    h1(doc, "Inputs reference", accent=INK, bookmark_name="qs_inputs")
    body(doc, "Column-level reference for the Inputs tab — the only tab you "
              "edit. Everything else in the workbook flows from here.",
         size=11, italic=True, color=MUTE)

    # --- Members ---------------------------------------------------------
    h2(doc, "Section 1 — Members (rows 6–12)")
    body(doc, "Up to four members. Edit columns A (label, free text) and B "
              "(TSB, dollars). Three derived columns are read-only:")
    bullets(doc, [
        ("Split %", "TSB ÷ total fund TSB."),
        ("Band 1 (3m–10m)", "MAX(0, MIN(TSB, $10m) − $3m) ÷ TSB. Taxed at 15%."),
        ("Band 2 (> 10m)", "MAX(0, TSB − $10m) ÷ TSB. Taxed at 25%."),
    ], tight=True)
    body(doc, "Traffic-light banner above the Members table reads green "
              "(no Div 296), amber ($3m–$10m), or deep amber (above $10m). Use "
              "it as the applicability check before reading any other number.")

    # --- Asset register ---------------------------------------------------
    h2(doc, "Section 2 — Asset register (rows 15–65, 50 rows)")
    simple_table(doc, ["Col", "Header", "Notes"], [
        ("A", "Asset code",
         "Short identifier. Free text."),
        ("B", "Asset name",
         "Plain description (e.g. ‘Commercial property’)."),
        ("C", "Original cost base",
         "Pre-reset cost base — used for ordinary CGT regardless of the reset election."),
        ("D", "Current market value",
         "Context only; not used in tax calcs. Useful sanity-check for reviewers."),
        ("E", "Market value at 30 Jun 2026",
         "Load-bearing. Becomes the Div 296 cost base under the reset scenario. "
         "Garbage in, garbage out."),
        ("F", "Valuation source / date",
         "Free text. Mirrored verbatim into the Notes valuation log for audit trail."),
        ("G", "Projected sale proceeds",
         "Expected cash on disposal. Blank = row excluded from totals (the amber "
         "tripwire fires)."),
        ("H", "Projected gain/loss",
         "Locked formula = proceeds − original cost base. Do not edit."),
        ("I", "Held > 12 months?",
         "Yes / No dropdown. Drives the 1/3 CGT discount under s115-100 ITAA 1997."),
    ], header_fill=INK)

    # --- Advanced assumptions --------------------------------------------
    h2(doc, "Section 3 — Advanced assumptions (rows 67–75)")
    body(doc, "Eight named-range cells that feed every downstream formula. "
              "Touch only if the ATO publishes a change. Defaults match the "
              "enacted law at the build date:")
    bullets(doc, [
        "rate_tier1 = 15% (slice $3m–$10m).",
        "rate_tier2 = 25% (slice above $10m).",
        "threshold_1 = $3,000,000 · threshold_2 = $10,000,000.",
        "discount_rate = 1/3 (s115-100 ITAA 1997).",
        "fund_cgt_rate = 15% (accumulation phase).",
        "indexation_increment_1 / _2 (forward-modelling only — not applied to Year 1).",
    ], tight=True)

    _page_break(doc)


# ============================================================================
# Page 7 — Analyser + Comparison
# ============================================================================

def analyser_and_comparison(doc) -> None:
    h1(doc, "Reading the Analyser", accent=SUB, bookmark_name="qs_analyser")

    h2(doc, "Fund summary (rows 6–13)")
    body(doc, "Side by side: If no reset · If elected · Difference. Bottom row "
              "(‘Total Div 296 tax’) is the headline. The Difference column is "
              "signed; negative (green) = reset saves tax, positive (red) = reset "
              "costs more tax.")

    h2(doc, "Per-asset analysis (rows 15–67)")
    body(doc, "Always reflects the elected-reset scenario. Most-read columns:")
    p = _para(doc, after=3)
    _add_bookmark(p, "qs_perasset", _next_bookmark_id())
    bullets(doc, [
        ("K — Div 296 tax",
         "Authoritative per-asset Div 296 tax (pro-rata of the elected-reset headline)."),
        ("L — Reset impact",
         "Signed delta. Negative = reset reduces this asset's Div 296 burden. "
         "Positive = reset-trap candidate."),
        ("E, G — Ord gross gain / Ord CGT (info only)",
         "Greyed. Do NOT sum to the fund Ord CGT because s102-5 capital-loss "
         "netting happens at the fund level, not per asset. Read the Reconciliation "
         "panel for the authoritative number."),
    ], tight=True)

    h2(doc, "Reconciliation (rows 70–74)")
    p_recon = _para(doc, after=3)
    _add_bookmark(p_recon, "qs_recon", _next_bookmark_id())
    _run(p_recon,
         "Authoritative fund-level totals: ordinary CGT after intra-year "
         "capital-loss netting (s102-5 method, then s115-100 1/3 discount on "
         "the discountable residue), Division 296 tax payable, and any "
         "current-year net unused capital loss carried forward.",
         size=10.5)

    # ---- Comparison tearsheet -------------------------------------------
    h1(doc, "Comparison tearsheet", accent=COPPER, bookmark_name="qs_comparison")
    body(doc, "Print-ready landscape A4. Use as a client meeting handout or "
              "file note without editing. Order top to bottom:")
    bullets(doc, [
        "Header strip — fund, prepared by, date, version.",
        "Members & TSB strip.",
        "Headline cards — Div 296 tax under each scenario plus signed net effect.",
        "Per-scenario subtotals — Earnings, Ord CGT (unchanged), Div 296 tax, "
        "Total burden.",
        "Per-member breakdown — TSB and Div 296 tax for both scenarios.",
        "Per-asset detail — top 10 by |Δ Div 296 tax|. Overflow note points "
        "at the Analyser for the full register.",
    ], tight=True)

    _page_break(doc)


# ============================================================================
# Page 8 — Caveats + Troubleshooting
# ============================================================================

def caveats_and_troubleshooting(doc) -> None:
    h1(doc, "Caveats", accent=AMBER, bookmark_name="qs_caveats")
    body(doc, "Hold these in mind whenever you read a figure. They are "
              "material and several work in opposing directions.",
         size=11, italic=True, color=MUTE)
    bullets(doc, [
        ("Prior-year carry-forward capital losses are NOT modelled.",
         "Apply any prior-year carry-forward outside the model."),
        ("Pension phase is NOT modelled.",
         "Assumes 100% accumulation phase. Overstates ordinary CGT for funds "
         "with pension members."),
        ("Reset-OFF scenario is realised-only.",
         "Under enacted Division 296, no-reset earnings are assessed on "
         "year-on-year TSB movement (realised + unrealised). Comparison "
         "understates the no-reset burden."),
        ("Multi-member splits are TSB-proportional.",
         "Real funds use an actuarial certificate on time-weighted balances."),
        ("Wash sale / Part IVA risk.",
         "Pre-30 June 2026 disposals undertaken purely for tax outcome sit in "
         "anti-avoidance territory (TR 2008/1, Part IVA)."),
        ("Transaction costs not modelled.",
         "Brokerage, stamp duty, liquidity drag, market-timing risk excluded."),
        ("Sheet protection is tamper-evident, not tamper-proof.",
         "Protection (where applied) is passwordless."),
    ], tight=True)

    h1(doc, "Troubleshooting", accent=WARN, bookmark_name="qs_troubleshooting")
    simple_table(doc, ["You see", "What it means", "Do this"], [
        ("Yellow ‘Sample data preloaded’ badge persists.",
         "Sample input (P1/S1/L1 codes or seeded TSBs) still present.",
         "Overwrite every sample cell on Inputs; badge clears automatically."),
        ("Red banner: proceeds without 30 Jun 2026 value.",
         "Those rows are excluded from every figure.",
         "Fill column E, or clear column G."),
        ("Red banner: formulas in the register.",
         "A non-values paste left live formulas in register cells.",
         "Press Ctrl+Z to undo, then Paste-Special > Values."),
        ("TSB diagnostic banner is deep amber.",
         "At least one member's TSB exceeds $10m — 25% above-$10m tier applies.",
         "Read Comparison tearsheet for the modelled impact. No input change."),
        ("Per-asset Ord CGT (col G) doesn't sum to Reconciliation Ord CGT.",
         "Correct behaviour. Loss-netting at fund level (s102-5), not per asset.",
         "Rely on Reconciliation for the authoritative figure."),
    ], header_fill=WARN)

    _page_break(doc)


# ============================================================================
# Page 9 — Cheat sheet (back cover)
# ============================================================================

def cheat_sheet(doc) -> None:
    h1(doc, "Cheat sheet", accent=INK, bookmark_name="qs_cheatsheet")
    body(doc, "Tear-off reference. Keep next to the workbook.",
         size=11, italic=True, color=MUTE)

    # 1. Inputs columns
    h2(doc, "Inputs columns at a glance")
    bullets(doc, [
        "A · Asset code (free text)",
        "B · Asset name (free text)",
        "C · Original cost base (ordinary CGT basis — unchanged by reset)",
        "D · Current market value (context only)",
        "E · Market value at 30 Jun 2026 (load-bearing — Div 296 basis under reset)",
        "F · Valuation source / date",
        "G · Projected sale proceeds (blank → row excluded)",
        "H · Projected gain/loss (locked formula — do not edit)",
        "I · Held > 12 months? (drives 1/3 CGT discount)",
    ], tight=True, size=10)

    # 2. Banners
    h2(doc, "Banners — colour, meaning, action")
    simple_table(doc, ["Colour", "Means", "Do this"], [
        ("Yellow", "Sample data still loaded.", "Overwrite every sample cell."),
        ("Red", "Formulas in register (Ctrl+V used).", "Ctrl+Z, Paste-Special > Values."),
        ("Red", "Proceeds without 30 Jun 2026 value.", "Fill col E or clear col G."),
        ("Amber", "TSB > $10m — 25% tier applies.", "No action; read Comparison."),
    ])

    # 3. Keyboard tricks
    h2(doc, "Keyboard tricks")
    bullets(doc, [
        ("Paste-Special > Values", "Ctrl+Alt+V then V — the only safe way to paste into the register."),
        ("Undo", "Ctrl+Z — recovers from a bad paste."),
        ("Word Navigation Pane", "Ctrl+F1 (this guide) → View → Navigation Pane."),
    ], tight=True, size=10)

    # 4. Where the headline lives
    h2(doc, "Where the headline lives")
    body(doc, "Comparison tab → Headline cards at the top. Three numbers: "
              "If no reset · If elected · Net effect (signed).", size=10)

    # 5. Five caveats that change numbers materially
    h2(doc, "Five caveats that change numbers")
    bullets(doc, [
        "Prior-year carry-forward losses — apply outside the model.",
        "Pension phase — not modelled (model overstates Ord CGT for pension members).",
        "No-reset is realised-only — understates no-reset Div 296 burden.",
        "Multi-member split is TSB-proportional, not actuarial.",
        "Transaction costs (brokerage, stamp duty, liquidity) excluded.",
    ], tight=True, size=10)

    # 6. Verifying the version
    h2(doc, "Verifying the version")
    body(doc, "Notes tab → unhide the rows immediately below the valuation "
              "log to see build_version, build_date, git_short_sha.", size=10)


# ============================================================================
# Continuous section breaks for per-section headers
# ============================================================================

def _add_per_section_headers(doc):
    """Insert a per-section header on each section: section name (left) +
    Adviser Edition v… (right). All section breaks are CONTINUOUS so no
    forced page breaks beyond our explicit ones."""
    # The doc starts with one section. We'll re-stamp its header to be
    # consistent. Adding more sections is left as a future enhancement —
    # multi-section headers in python-docx are fiddly, and the running
    # footer from the workbook's _stamp_print_footer pattern isn't used
    # here (the guide is a different document). One uniform header keeps
    # the impl simple and the visual consistent.
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p, f"Adviser Edition  ·  v{__version__}", size=8, italic=True, color=MUTE)


# ============================================================================
# Entry point
# ============================================================================

def build_guide() -> Document:
    doc = Document()
    _set_margins(doc)
    _add_per_section_headers(doc)

    cover(doc)
    quick_start(doc)
    workbook_map(doc)
    contents_page(doc)
    inputs_reference(doc)
    analyser_and_comparison(doc)
    caveats_and_troubleshooting(doc)
    cheat_sheet(doc)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_guide()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
